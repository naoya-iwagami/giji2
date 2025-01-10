import os    
import io    
import docx    
from docx import Document    
import streamlit as st    
from openai import AzureOpenAI    
from reportlab.pdfgen import canvas    
from reportlab.lib.pagesizes import A4    
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable    
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle    
  
# クリップボード操作のためのライブラリをインポート    
import pyperclip    
  
# ページの設定をwideにする    
st.set_page_config(layout="wide")    
  
# 環境変数からAPIキーを取得    
api_key = os.getenv("OPENAI_API_KEY")    
if not api_key:    
    st.error("APIキーが設定されていません。環境変数OPENAI_API_KEYを設定してください。")    
    st.stop()    
  
# Azure OpenAI クライアントの設定    
client = AzureOpenAI(    
    api_key=api_key,    
    api_version="2024-10-01-preview",    
    azure_endpoint="https://test-chatgpt-pm-1.openai.azure.com/"  # あなたのエンドポイントに置き換えてください    
)    
  
def read_docx(file):    
    try:    
        doc = docx.Document(file)    
        full_text = [para.text for para in doc.paragraphs]    
        return '\n'.join(full_text)    
    except Exception as e:    
        st.error(f"Error reading docx file: {e}")    
        return None    
  
def create_summary(full_text, system_message):    
    try:    
        # ユーザーメッセージを作成    
        user_message = f"{system_message}\n\n以下はテキストです。\n\n{full_text}"    
  
        messages = [    
            {"role": "user", "content": user_message}    
        ]    
  
        response = client.chat.completions.create(    
            model='o1-preview',    
            messages=messages,    
            max_completion_tokens=32768  # 最大トークン数を指定    
        )    
  
        final_summary = response.choices[0].message.content    
        return final_summary    
  
    except Exception as e:    
        st.error(f"Unexpected error in create_summary: {e}")    
        return None    
  
# テキストを解析してフォーマット情報を抽出する関数    
def parse_text(text):    
    lines = text.split('\n')    
    parsed_content = []    
    for line in lines:    
        line = line.strip()    
        if line.startswith('### '):    
            parsed_content.append(('heading3', line[4:]))    
        elif line.startswith('## '):    
            parsed_content.append(('heading2', line[3:]))    
        elif line.startswith('# '):    
            parsed_content.append(('heading1', line[2:]))    
        elif line.startswith('- '):    
            parsed_content.append(('list', line[2:]))    
        elif line == '':    
            parsed_content.append(('newline', ''))    
        else:    
            parsed_content.append(('paragraph', line))    
    return parsed_content    
  
# Word文書を作成する関数    
def create_word_document(parsed_content):    
    doc = Document()    
    for content_type, content in parsed_content:    
        if content_type == 'heading1':    
            doc.add_heading(content, level=1)    
        elif content_type == 'heading2':    
            doc.add_heading(content, level=2)    
        elif content_type == 'heading3':    
            doc.add_heading(content, level=3)    
        elif content_type == 'list':    
            p = doc.add_paragraph(style='List Bullet')    
            p.add_run(content)    
        elif content_type == 'paragraph':    
            doc.add_paragraph(content)    
        elif content_type == 'newline':    
            doc.add_paragraph()    
    return doc    
  
# PDFを作成する関数    
def create_pdf(parsed_content):    
    pdf_stream = io.BytesIO()    
    doc = SimpleDocTemplate(pdf_stream, pagesize=A4)    
    styles = getSampleStyleSheet()    
  
    # 日本語フォントの登録    
    from reportlab.pdfbase import pdfmetrics    
    from reportlab.pdfbase.ttfonts import TTFont    
  
    # フォントファイルのパスを指定（'ipaexg.ttf' をスクリプトと同じディレクトリに配置する）    
    pdfmetrics.registerFont(TTFont('IPAexGothic', 'ipaexg.ttf'))    
  
    # カスタムスタイルの定義（フォントに 'IPAexGothic' を指定）    
    styles.add(ParagraphStyle(name='CustomHeading1', parent=styles['Heading1'],    
                              fontName='IPAexGothic', fontSize=16, leading=20))    
    styles.add(ParagraphStyle(name='CustomHeading2', parent=styles['Heading2'],    
                              fontName='IPAexGothic', fontSize=14, leading=18))    
    styles.add(ParagraphStyle(name='CustomHeading3', parent=styles['Heading3'],    
                              fontName='IPAexGothic', fontSize=12, leading=16))    
    styles.add(ParagraphStyle(name='CustomNormal', parent=styles['Normal'],    
                              fontName='IPAexGothic', fontSize=10, leading=14))    
  
    flowables = []    
  
    bullet_items = []    
  
    for content_type, content in parsed_content:    
        if content_type == 'heading1':    
            flowables.append(Paragraph(content, styles['CustomHeading1']))    
            flowables.append(Spacer(1, 12))    
        elif content_type == 'heading2':    
            flowables.append(Paragraph(content, styles['CustomHeading2']))    
            flowables.append(Spacer(1, 10))    
        elif content_type == 'heading3':    
            flowables.append(Paragraph(content, styles['CustomHeading3']))    
            flowables.append(Spacer(1, 8))    
        elif content_type == 'list':    
            bullet_items.append(Paragraph(content, styles['CustomNormal']))    
        elif content_type == 'paragraph':    
            if bullet_items:    
                flowables.append(ListFlowable(    
                    bullet_items,    
                    bulletType='bullet',    
                    bulletFontName='IPAexGothic',  # バレットにもフォントを適用    
                ))    
                bullet_items = []    
            flowables.append(Paragraph(content, styles['CustomNormal']))    
            flowables.append(Spacer(1, 6))    
        elif content_type == 'newline':    
            if bullet_items:    
                flowables.append(ListFlowable(    
                    bullet_items,    
                    bulletType='bullet',    
                    bulletFontName='IPAexGothic',    
                ))    
                bullet_items = []    
            flowables.append(Spacer(1, 12))    
  
    # 最後に残った箇条書きを追加    
    if bullet_items:    
        flowables.append(ListFlowable(    
            bullet_items,    
            bulletType='bullet',    
            bulletFontName='IPAexGothic',    
        ))    
  
    doc.build(flowables)    
    pdf_stream.seek(0)    
    return pdf_stream    
  
# StreamlitのUI設定    
st.title("議事録アプリ")    
st.write("docxファイルをアップロードして、その内容から議事録を作成します。")    
  
# ユーザーが指定できる system_message（デフォルトの文章を設定）    
default_system_message = """あなたはプロフェッショナルな議事録作成者です。    
以下の要件に従って、添付したテキストを基に、できるだけ詳細かつ具体的な議事録を作成して内容を咀嚼しプロットしてください。    
- 議事録はMarkdown形式で作成してください。      
- 以下のような専門用語が使われている可能性があるので文章修正の参考にしてください。    
    - 例えば、讃岐という言葉が製品名を指してしそうな文脈で出てきたら、SANUQIという文字に修正してください。    
###        
工場名: L1, L2, L3, L4, L5, L6, L7, L8, L9, LX (LX1), LX2, M10,C1        
サイト名=事業所名(設置工場名) : 神戸(L1,L2,L5,L6,C1)、神戸第2(L7,L8)、西神(L3,L4)、甲府(LX1,LX2)、日野(M10)、TB(L9)        
製品名: 3XR, 4UA, 2UA, SANUQI(さぬきと読む), SAZMA(さつまと読む), 4CT, 6VQ, 4VQ, 4TQ, 6TQ, SWPL        
製品の幅の名前: W (ワイドと読む), SW(スーパーワイドと読む), UW(ウルトラワイドと読む), EUW, EXW, UWD        
部署名: 生技(生産技術), 生管(生産管理), 品証(品質保証), 設備(設備保全グループ), 第1製造, 第2製造, 第3製造, 関西生産, 関東生産        
顧客名：CSOT, BOE, SDP, LGD, AUO, イノラックス, SDC, CHOT, HMO, SDI, DNP, SAPO, SNP, WINPOL, CMMT, 凸版, 杉金, 住友化学, 日本製紙, 東友, 日東        
競合名：Zeon（ゼオンと読む）, Fuji, 東洋紡        
サプライヤー名：ECC, JSR, ダイセル        
材料名：TAC(タックと読む), CAP（キャップと読む）, ARTON（アートンと読む）    
"""    
  
# セッション状態の初期化    
if 'summary' not in st.session_state:    
    st.session_state['summary'] = None    
if 'edited_summary' not in st.session_state:    
    st.session_state['edited_summary'] = None    
  
# ユーザー入力をフォームで受け取る    
with st.form(key='input_form'):    
    system_message = st.text_area("システムメッセージを入力してください:", default_system_message, height=300)    
        
    # 複数ファイルのアップロード    
    uploaded_files = st.file_uploader(    
        "ファイルをアップロード (最大3つ)",    
        type="docx",    
        accept_multiple_files=True    
    )    
        
    # 送信ボタン    
    submit_button = st.form_submit_button(label='議事録を作成')    
  
if submit_button:    
    if uploaded_files:    
        if len(uploaded_files) > 3:    
            st.error("最大3つのファイルまでアップロードできます。")    
        else:    
            total_char_count = 0    
            file_char_counts = []    
            all_texts = ""  # 全文テキストを保存する変数    
  
            for uploaded_file in uploaded_files:    
                text = read_docx(uploaded_file)    
                if text:    
                    char_count = len(text)    
                    total_char_count += char_count    
                    file_char_counts.append((uploaded_file.name, char_count))    
                    all_texts += text + "\n"    
                else:    
                    st.error(f"{uploaded_file.name} の読み込みに失敗しました。")    
  
            # 各ファイルの文字数と合計文字数を表示    
            st.write("アップロードされたファイルの文字数:")    
            for file_name, char_count in file_char_counts:    
                st.write(f"- {file_name}: {char_count}文字")    
            st.write(f"**合計文字数: {total_char_count}文字**")    
  
            if all_texts:    
                # 議事録を作成    
                summary = create_summary(all_texts, system_message)    
  
                if summary is None:    
                    st.error("議事録の生成に失敗しました。")    
                else:    
                    st.session_state['summary'] = summary  # セッション状態に保存    
                    st.session_state['edited_summary'] = summary  # 編集用のテキストも初期化    
                    st.success("議事録が作成されました。内容を編集できます。")    
            else:    
                st.error("ファイルの読み込みに失敗しました。")    
    else:    
        st.error("ファイルをアップロードしてください。")    
  
# 議事録が生成されている場合、編集とダウンロードを表示    
if st.session_state['summary']:    
    st.header("作成された議事録（編集可能）")    
  
    # 編集可能なテキストエリアを表示    
    edited_summary = st.text_area(    
        "議事録を編集してください:",    
        st.session_state['edited_summary'],    
        height=400    
    )    
  
    # 編集内容をセッション状態に保存    
    st.session_state['edited_summary'] = edited_summary    
  
    # 解析済みテキストを取得    
    parsed_content = parse_text(st.session_state['edited_summary'])    
  
    # Wordドキュメントの作成    
    doc = create_word_document(parsed_content)    
    docx_stream = io.BytesIO()    
    doc.save(docx_stream)    
    docx_stream.seek(0)    
  
    # PDFの作成    
    pdf_stream = create_pdf(parsed_content)    
  
    # ボタンを縦一列に配置    
    if st.button("議事録コピー"):    
        pyperclip.copy(st.session_state['edited_summary'])    
        st.success("議事録の内容がクリップボードにコピーされました。Outlookで新規メールを開き、本文に貼り付けてください。")    
  
    # PDFダウンロードボタンの表示    
    st.download_button(    
        label="議事録ダウンロード（PDF）",    
        data=pdf_stream,    
        file_name='議事録.pdf',    
        mime='application/pdf'    
    )    
  
    # Wordダウンロードボタンの表示    
    st.download_button(    
        label="議事録ダウンロード（Word）",    
        data=docx_stream,    
        file_name='議事録.docx',    
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'    
    )  